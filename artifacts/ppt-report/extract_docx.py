from pathlib import Path
import sys

from docx import Document


def main() -> None:
    path = Path(sys.argv[1])
    doc = Document(path)
    print(f"FILE: {path}")
    print(f"PARAGRAPHS: {len(doc.paragraphs)}")
    print(f"TABLES: {len(doc.tables)}")
    print(f"INLINE_SHAPES: {len(doc.inline_shapes)}")
    print("\n=== PARAGRAPHS ===")
    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            print(f"P{idx:03d} [{paragraph.style.name}]: {text}")
    for table_idx, table in enumerate(doc.tables, start=1):
        print(f"\n=== TABLE {table_idx} ({len(table.rows)}x{len(table.columns)}) ===")
        for row_idx, row in enumerate(table.rows, start=1):
            values = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
            print(f"R{row_idx:03d}: " + " | ".join(values))


if __name__ == "__main__":
    main()
