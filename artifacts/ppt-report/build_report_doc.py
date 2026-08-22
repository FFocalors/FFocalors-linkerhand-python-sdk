from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SOURCE = Path(__file__).with_name("report_source.md")
OUT = ROOT / "LinkerHand_项目汇报PPT母稿_已实现功能与实现方法.docx"
FONT = "Microsoft YaHei"
BLUE, DARK_BLUE, NAVY = "2E74B5", "1F4D78", "18324A"
INK, MUTED, LIGHT_BLUE, LIGHT_GRAY = "1F2937", "64748B", "E8F1F8", "F2F4F7"


def font(run, size=11, color=INK, bold=False, italic=False, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia"):
        rpr.rFonts.set(qn(f"w:{key}"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold, run.italic = bold, italic


def shade(target, fill):
    pr = target._p.get_or_add_pPr() if hasattr(target, "_p") else target._tc.get_or_add_tcPr()
    node = pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        pr.append(node)
    node.set(qn("w:fill"), fill)


def paragraph_border_left(paragraph, color):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    pbdr.append(left)


def add_num_defs(doc):
    numbering = doc.part.numbering_part.element
    abs_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]

    def add(fmt, marker):
        abs_id = max(abs_ids + [0]) + 1
        abs_ids.append(abs_id)
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        for tag, value in (("start", "1"), ("numFmt", fmt), ("lvlText", marker), ("suff", "tab")):
            el = OxmlElement(f"w:{tag}")
            el.set(qn("w:val"), value)
            lvl.append(el)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, ind, spacing])
        lvl.append(ppr)
        abstract.append(lvl)
        numbering.append(abstract)
        num_id = max(num_ids + [0]) + 1
        num_ids.append(num_id)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abs_id))
        num.append(ref)
        numbering.append(num)
        return num_id

    return add("bullet", "•"), add("decimal", "%1.")


def apply_num(paragraph, num_id):
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, nid])
    paragraph._p.get_or_add_pPr().append(numpr)


def setup(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name, normal.font.size = FONT, Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        s = styles[name]
        s.font.name, s.font.size, s.font.bold = FONT, Pt(size), True
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    hp = sec.header.paragraphs[0]
    font(hp.add_run("LinkerHand 机械手 GUI 控制系统 | 项目汇报 PPT 母稿"), 8.5, MUTED)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(fp.add_run("第 "), 9, MUTED)
    run = fp.add_run()
    begin, instr, end = OxmlElement("w:fldChar"), OxmlElement("w:instrText"), OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    font(fp.add_run(" 页"), 9, MUTED)
    return add_num_defs(doc)


def set_table_geometry(table, widths):
    total = sum(widths)
    pr = table._tbl.tblPr
    for tag, value in (("tblW", total), ("tblInd", 120)):
        node = pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW")) or OxmlElement("w:tcW")
            if tcw.getparent() is None:
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")
            mar = OxmlElement("w:tcMar")
            for side, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"), str(value))
                m.set(qn("w:type"), "dxa")
                mar.append(m)
            tcpr.append(mar)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows):
    cols = len(rows[0])
    if cols == 2:
        widths = [3000, 6360]
    elif cols == 3:
        widths = [1700, 3500, 4160]
    else:
        widths = [9360 // cols] * cols
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for r_idx, values in enumerate(rows):
        cells = table.rows[0].cells if r_idx == 0 else table.add_row().cells
        for c_idx, value in enumerate(values):
            if r_idx == 0:
                shade(cells[c_idx], LIGHT_GRAY)
            elif r_idx % 2 == 0:
                shade(cells[c_idx], "F7F9FC")
            p = cells[c_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            font(p.add_run(value), 9.0, NAVY if r_idx == 0 else INK, r_idx == 0)
    trpr = table.rows[0]._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    trpr.append(marker)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_text(doc, text, size=11, color=INK, bold=False, italic=False, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    font(p.add_run(text), size, color, bold, italic)
    return p


def parse_source(doc, bullet_id, number_id):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code, code_lines, table_rows = False, [], []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            flush_table()
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.right_indent = Inches(0.12)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.05
                shade(p, "F3F6F8")
                paragraph_border_left(p, DARK_BLUE)
                font(p.add_run("\n".join(code_lines)), 9.3, DARK_BLUE, name="Consolas")
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.fullmatch(r"[-: ]+", c) for c in cells):
                continue
            table_rows.append(cells)
            continue
        flush_table()
        if not line:
            continue
        if line == "---PAGE---":
            doc.add_page_break()
        elif line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.space_before = Pt(0)
            font(p.add_run(line[2:]), 16, BLUE, True)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            font(p.add_run(line[3:]), 13, BLUE, True)
        elif line.startswith("> "):
            body = line[2:]
            label, sep, rest = body.partition("：")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.right_indent = Inches(0.08)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.15
            shade(p, LIGHT_BLUE if label != "证据口径" else "F8FAFC")
            paragraph_border_left(p, BLUE if label != "证据口径" else MUTED)
            font(p.add_run((label + "：") if sep else ""), 11, BLUE, True)
            font(p.add_run(rest if sep else body), 11, INK)
        elif line.startswith("- "):
            p = doc.add_paragraph()
            apply_num(p, bullet_id)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            font(p.add_run(line[2:]), 11, INK)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph()
            apply_num(p, number_id)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            font(p.add_run(re.sub(r"^\d+\. ", "", line)), 11, INK)
        elif line.startswith("[[IMAGE:"):
            payload = line[len("[[IMAGE:"):-2]
            rel, caption, width = (payload.split("|") + ["", "5.4"])[:3]
            path = ROOT / rel
            if path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(3)
                p.add_run().add_picture(str(path), width=Inches(float(width)))
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(7)
                font(cp.add_run(caption), 9, MUTED, italic=True)
        else:
            add_text(doc, line)
    flush_table()


def build():
    doc = Document()
    bullet_id, number_id = setup(doc)
    parse_source(doc, bullet_id, number_id)
    doc.core_properties.title = "LinkerHand 项目汇报 PPT 母稿：已实现功能与实现方法"
    doc.core_properties.subject = "LinkerHand Python SDK 二次开发项目汇报"
    doc.core_properties.author = "LinkerHand 项目组"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
