from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
PATH = ROOT / "灵伴巧手_具身智能家庭辅助机械手实践报告.docx"


def clean_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = run.text.replace("**", "").replace("`", "")


doc = Document(PATH)
for paragraph in doc.paragraphs:
    clean_paragraph(paragraph)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                clean_paragraph(paragraph)
doc.save(PATH)
print(PATH)
